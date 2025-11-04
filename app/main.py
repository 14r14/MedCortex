"""Main Streamlit application for MedCortex AI Research Analyst.

This module provides the user interface for the MedCortex application,
including document upload, query processing, and report generation.
"""

import base64
import hashlib
import re
import uuid
from datetime import datetime
from typing import List, Optional, Tuple

import streamlit as st
from dotenv import load_dotenv

from app.config import Settings
from app.rag.pipeline import IngestionPipeline, QueryPipeline


@st.cache_resource
def get_css_content() -> str:
    """Return CSS content for MedCortex UI styling.

    Returns:
        CSS content as a string. Cached to avoid re-injecting on every rerun.
    """
    return r"""
        <style>
        /* MedCortex Color Palette */
        
        /* Light Mode Colors */
        :root {
            --background: #FFFFFF;
            --ui-panel: #F0F2F6;
            --primary-text: #121212;
            --secondary-text: #525252;
            --verification-green: #22C55E;
            --warning-orange: #F97316;
            --user-error-red: #EF4444;
            --info-blue: #3B82F6;
        }
        
        /* Dark Mode Colors */
        @media (prefers-color-scheme: dark) {
            :root {
                --background: #121212;
                --ui-panel: #2b2b2b;
                --primary-text: #FFFFFF;
                --secondary-text: #AAAAAA;
                --verification-green: #22C55E;
                --warning-orange: #F97316;
                --user-error-red: #EF4444;
                --info-blue: #3B82F6;
            }
        }
        
        /* Override Streamlit default styles */
        .stApp {
            background-color: var(--background);
        }
        
        /* Main container */
        .main .block-container {
            background-color: var(--background);
            padding-top: 2rem;
            padding-bottom: 2rem;
        }
        
        /* Title styling - MedCortex branding */
        h1 {
            color: var(--primary-text) !important;
            font-weight: 700;
            font-size: 2.5rem;
            letter-spacing: -0.02em;
        }
        
        /* MedCortex logo/title accent - removed icon for professional look */
        
        /* Caption/subtitle */
        .stMarkdown p {
            color: var(--primary-text);
        }
        
        /* Sidebar */
        .css-1d391kg {
            background-color: var(--ui-panel);
        }
        
        [data-testid="stSidebar"] {
            background-color: var(--ui-panel);
        }
        
        /* Sidebar title - MedCortex */
        [data-testid="stSidebar"] h1 {
            color: var(--primary-text) !important;
            font-weight: 700;
            font-size: 1.75rem;
            letter-spacing: -0.01em;
            margin-bottom: 0.5rem !important;
        }
        
        [data-testid="stSidebar"] .stMarkdown h2,
        [data-testid="stSidebar"] .stMarkdown h3 {
            color: var(--primary-text) !important;
        }
        
        [data-testid="stSidebar"] .stMarkdown,
        [data-testid="stSidebar"] p {
            color: var(--secondary-text) !important;
        }
        
        /* Sidebar dividers */
        [data-testid="stSidebar"] hr {
            border-color: var(--secondary-text) !important;
            opacity: 0.3 !important;
            margin: 1rem 0 !important;
        }
        
        /* Sidebar document items */
        [data-testid="stSidebar"] .stMarkdown strong {
            color: var(--primary-text) !important;
            font-weight: 600;
        }
        
        [data-testid="stSidebar"] .stMarkdown .stCaption {
            color: var(--secondary-text) !important;
            font-size: 0.875rem;
        }
        
        /* Sidebar info box */
        [data-testid="stSidebar"] .stInfo {
            background-color: rgba(59, 130, 246, 0.1) !important;
            border-left: 3px solid var(--info-blue) !important;
        }
        
        /* User chat message - User/Error Red */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="user"]) [data-testid="stChatMessageContent"] {
            background-color: var(--user-error-red);
            color: #ffffff;
            border-radius: 12px 12px 0 12px;
            padding: 12px 16px;
        }
        
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="user"]) [data-testid="stChatMessageContent"] p {
            color: #ffffff !important;
        }
        
        /* Assistant chat message - UI Panel */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] {
            background-color: var(--ui-panel);
            color: var(--primary-text) !important;
            border-radius: 12px 12px 12px 0;
            padding: 12px 16px;
        }
        
        /* Ensure all text in assistant messages is primary text - comprehensive selector */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] * {
            color: var(--primary-text) !important;
        }
        
        /* Override markdown text colors within assistant messages */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] .stMarkdown,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] .stMarkdown *,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] p,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] div,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] span,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] li,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] ul,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] ol,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] h1,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] h2,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] h3,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] h4,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] h5,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] h6,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] code,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] pre,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] blockquote,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] em,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] i,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] small {
            color: var(--primary-text) !important;
        }
        
        /* Bold/strong text in assistant messages */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] strong,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] b {
            color: var(--primary-text) !important;
        }
        
        /* Links in assistant messages */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] a {
            color: var(--primary-text) !important;
            text-decoration: underline !important;
        }
        
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] a:hover {
            color: var(--primary-text) !important;
            opacity: 0.9 !important;
        }
        
        /* ============ CURSOR-STYLE BUTTONS ============ */
        /* Base Cursor-style button - applies to all buttons by default */
        .stButton > button,
        [data-testid="stButton"] > button,
        .stDownloadButton > button,
        button[data-testid="baseButton-secondary"],
        button[data-testid="baseButton-primary"] {
            /* Cursor-style base appearance */
            background-color: var(--ui-panel) !important;
            color: var(--primary-text) !important;
            border: 1px solid rgba(82, 82, 82, 0.3) !important;
            border-radius: 6px !important;
            padding: 0.375rem 0.75rem !important;
            font-size: 0.8125rem !important;
            font-weight: 400 !important;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif !important;
            cursor: pointer !important;
            transition: all 0.15s ease !important;
            box-shadow: none !important;
            outline: none !important;
            min-height: auto !important;
            height: auto !important;
            white-space: nowrap !important;
        }
        
        /* Hover state */
        .stButton > button:hover,
        [data-testid="stButton"] > button:hover,
        .stDownloadButton > button:hover,
        button[data-testid="baseButton-secondary"]:hover,
        button[data-testid="baseButton-primary"]:hover {
            background-color: var(--ui-panel) !important;
            opacity: 0.9 !important;
            border-color: var(--secondary-text) !important;
        }
        
        /* Active/pressed state - no special styling */
        .stButton > button:active,
        [data-testid="stButton"] > button:active,
        .stDownloadButton > button:active,
        button[data-testid="baseButton-secondary"]:active,
        button[data-testid="baseButton-primary"]:active {
            /* No active styles - keep same as default */
        }
        
        /* Focus state - no special styling */
        .stButton > button:focus,
        [data-testid="stButton"] > button:focus,
        .stDownloadButton > button:focus,
        button[data-testid="baseButton-secondary"]:focus,
        button[data-testid="baseButton-primary"]:focus,
        .stButton > button:focus-visible,
        [data-testid="stButton"] > button:focus-visible,
        .stDownloadButton > button:focus-visible,
        button[data-testid="baseButton-secondary"]:focus-visible,
        button[data-testid="baseButton-primary"]:focus-visible {
            /* No focus styles - keep same as default */
            outline: none !important;
        }
        
        /* Disabled state */
        .stButton > button:disabled,
        [data-testid="stButton"] > button:disabled,
        .stDownloadButton > button:disabled,
        button[data-testid="baseButton-secondary"]:disabled,
        button[data-testid="baseButton-primary"]:disabled {
            opacity: 0.5 !important;
            cursor: not-allowed !important;
        }
        
        /* Dark mode support */
        @media (prefers-color-scheme: dark) {
            .stButton > button,
            [data-testid="stButton"] > button,
            .stDownloadButton > button,
            button[data-testid="baseButton-secondary"],
            button[data-testid="baseButton-primary"] {
                background-color: var(--ui-panel) !important;
                color: var(--primary-text) !important;
                border-color: rgba(170, 170, 170, 0.3) !important;
            }
            
            .stButton > button:hover,
            [data-testid="stButton"] > button:hover,
            .stDownloadButton > button:hover,
            button[data-testid="baseButton-secondary"]:hover,
            button[data-testid="baseButton-primary"]:hover {
                background-color: var(--ui-panel) !important;
                opacity: 0.9 !important;
                border-color: var(--secondary-text) !important;
            }
            
            .stButton > button:active,
            [data-testid="stButton"] > button:active,
            .stDownloadButton > button:active,
            button[data-testid="baseButton-secondary"]:active,
            button[data-testid="baseButton-primary"]:active {
                /* No active styles - keep same as default */
            }
        }
        
        /* Context-specific button styles - different styles for different locations */
        
        /* Main content buttons */
        .main .stButton > button {
            padding: 0.5rem 1rem !important;
            font-size: 0.875rem !important;
        }
        
        /* Chat area buttons (Add to Report) - smaller */
        [data-testid="stChatMessage"] ~ .stButton > button,
        .main [data-testid="stChatMessage"] ~ .stButton > button {
            padding: 0.375rem 0.75rem !important;
            font-size: 0.75rem !important;
        }
        
        /* Download buttons - same as base */
        .stDownloadButton > button {
            /* Uses base Cursor style */
        }
        
        /* ============ NAVIGATION SECTION (Sidebar) ============ */
        /* Reduce spacing after Navigation heading */
        [data-testid="stSidebar"] h3,
        [data-testid="stSidebar"] .stMarkdown h3 {
            margin-bottom: 1rem !important;
            margin-top: 0 !important;
            padding-bottom: 0 !important;
            padding-top: 0 !important;
            line-height: 1.2 !important;
        }
        
        /* Reduce spacing between Navigation heading and buttons */
        [data-testid="stSidebar"] [data-testid="stVerticalBlock"]:has(h3) + [data-testid="stVerticalBlock"],
        [data-testid="stSidebar"] .stMarkdown:has(h3) ~ [data-testid="stVerticalBlock"] {
            margin-top: 0.1rem !important;
            padding-top: 0 !important;
        }
        
        /* Reduce spacing on button containers that follow Navigation heading */
        [data-testid="stSidebar"] [data-testid="stVerticalBlock"]:has(h3) ~ [data-testid="stVerticalBlock"] .stButton,
        [data-testid="stSidebar"] [data-testid="stVerticalBlock"]:has(h3) + [data-testid="stVerticalBlock"] .stButton {
            margin-top: 0 !important;
            padding-top: 0 !important;
            margin-bottom: 0.25rem !important;
        }
        
        /* Target the first button after Navigation more directly */
        [data-testid="stSidebar"] h3 ~ [data-testid="stVerticalBlock"] .stButton,
        [data-testid="stSidebar"] h3 + [data-testid="stVerticalBlock"] .stButton,
        [data-testid="stSidebar"] .stMarkdown:has(h3) + [data-testid="stVerticalBlock"] .stButton {
            margin-top: 0 !important;
            padding-top: 0 !important;
        }
        
        /* Navigation buttons in sidebar */
        [data-testid="stSidebar"] .stButton > button,
        [data-testid="stSidebar"] button[data-testid="baseButton-secondary"] {
            background-color: var(--ui-panel) !important;
            color: var(--primary-text) !important;
            border: 1px solid rgba(82, 82, 82, 0.3) !important;
            font-weight: 500 !important;
        }
        
        /* Ensure button text elements use primary color */
        [data-testid="stSidebar"] .stButton > button *,
        [data-testid="stSidebar"] button[data-testid="baseButton-secondary"] *,
        [data-testid="stSidebar"] .stButton > button span,
        [data-testid="stSidebar"] button[data-testid="baseButton-secondary"] span,
        [data-testid="stSidebar"] .stButton > button p,
        [data-testid="stSidebar"] button[data-testid="baseButton-secondary"] p {
            color: var(--primary-text) !important;
        }
        
        
        [data-testid="stSidebar"] .stButton > button:hover,
        [data-testid="stSidebar"] button[data-testid="baseButton-secondary"]:hover {
            background-color: var(--ui-panel) !important;
            opacity: 0.9 !important;
            border-color: var(--secondary-text) !important;
            color: var(--primary-text) !important;
        }
        
        [data-testid="stSidebar"] .stButton > button:hover *,
        [data-testid="stSidebar"] button[data-testid="baseButton-secondary"]:hover * {
            color: var(--primary-text) !important;
        }
        
        [data-testid="stSidebar"] .stButton > button:disabled,
        [data-testid="stSidebar"] button[data-testid="baseButton-secondary"]:disabled {
            background-color: var(--ui-panel) !important;
            opacity: 0.5 !important;
            color: var(--secondary-text) !important;
            border-color: rgba(82, 82, 82, 0.2) !important;
        }
        
        [data-testid="stSidebar"] .stButton > button:disabled *,
        [data-testid="stSidebar"] button[data-testid="baseButton-secondary"]:disabled * {
            color: var(--secondary-text) !important;
        }
        
        /* Headers with dividers - compact */
        h2 {
            color: var(--primary-text) !important;
            font-weight: 700 !important;
            margin-top: 1rem !important;
            margin-bottom: 0.75rem !important;
            font-size: 1.5rem !important;
        }
        
        /* Horizontal dividers - compact */
        hr {
            border-color: var(--info-blue) !important;
            border-width: 2px !important;
            margin: 1rem 0 !important;
        }
        
        /* Info messages */
        .stInfo {
            background-color: rgba(59, 130, 246, 0.1) !important;
            border-left: 4px solid var(--info-blue) !important;
            padding: 1rem !important;
            border-radius: 6px !important;
        }
        
        /* Status update text with pulsing animation */
        .status-update-text {
            color: var(--primary-text) !important;
            font-size: 0.95rem !important;
            padding: 0 !important;
            margin: 0 !important;
            animation: pulse 2s ease-in-out infinite;
            line-height: 1.5 !important;
        }
        
        /* Remove paragraph default margins for status text */
        [data-testid="stChatMessage"] [data-testid="stChatMessageContent"] p.status-update-text {
            margin: 0 !important;
            padding: 0 !important;
        }
        
        /* Align status container with chat message avatar - assistant messages */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] .status-update-text {
            color: var(--primary-text) !important;
            display: block !important;
            line-height: 1.5 !important;
        }
        
        /* Ensure status container has no extra padding */
        [data-testid="stChatMessage"] [data-testid="stChatMessageContent"] [data-testid="stVerticalBlock"]:has(.status-update-text),
        [data-testid="stChatMessage"] [data-testid="stChatMessageContent"] div:has(.status-update-text) {
            padding: 0 !important;
            margin: 0 !important;
        }
        
        @keyframes pulse {
            0%, 100% {
                opacity: 1;
            }
            50% {
                opacity: 0.6;
            }
        }
        
        /* Minimalistic expander styling - match status text style */
        [data-testid="stExpander"] {
            border: none !important;
            box-shadow: none !important;
            background: transparent !important;
            margin: 0.5rem 0 !important;
            outline: none !important;
        }
        
        [data-testid="stExpander"] > div {
            border: none !important;
            background: transparent !important;
            outline: none !important;
        }
        
        /* Expander header (button) - minimalistic, match status text */
        [data-testid="stExpander"] summary {
            border: none !important;
            outline: none !important;
            background: transparent !important;
            padding: 0 !important;
            margin: 0 !important;
            font-weight: 400 !important;
            font-size: 0.95rem !important;
            color: var(--primary-text) !important;
            cursor: pointer !important;
            transition: all 0.2s ease !important;
            border-radius: 8px !important;
            display: flex !important;
            align-items: center !important;
            list-style: none !important;
        }
        
        /* Expander in assistant chat messages */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stExpander"] summary {
            color: var(--primary-text) !important;
        }
        
        /* Curved hover effect - subtle background with rounded corners */
        [data-testid="stExpander"] summary:hover {
            background: rgba(0, 0, 0, 0.05) !important;
            border-radius: 8px !important;
        }
        
        /* Expander hover in assistant messages */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stExpander"] summary:hover {
            background: rgba(240, 240, 240, 0.1) !important;
        }
        
        /* Expander content area - distinguished with subtle background and indentation */
        [data-testid="stExpander"] [data-testid="stExpanderContent"] {
            border: none !important;
            outline: none !important;
            background: rgba(0, 0, 0, 0.02) !important;
            padding: 0.75rem 1rem !important;
            margin-top: 0.25rem !important;
            margin-left: 0.5rem !important;
            border-radius: 8px !important;
            border-left: 2px solid rgba(0, 0, 0, 0.1) !important;
        }
        
        /* For assistant chat messages, use lighter background and soft white border */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stExpander"] [data-testid="stExpanderContent"] {
            background: rgba(240, 240, 240, 0.05) !important;
            border-left: 2px solid rgba(240, 240, 240, 0.2) !important;
        }
        
        [data-testid="stExpander"][open] [data-testid="stExpanderContent"] {
            background: rgba(0, 0, 0, 0.03) !important;
            border-left: 2px solid rgba(0, 0, 0, 0.15) !important;
        }
        
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stExpander"][open] [data-testid="stExpanderContent"] {
            background: rgba(240, 240, 240, 0.08) !important;
            border-left: 2px solid rgba(240, 240, 240, 0.3) !important;
        }
        
        /* Ensure no borders on expander container */
        [data-testid="stExpander"] details {
            border: none !important;
            outline: none !important;
            background: transparent !important;
        }
        
        /* File uploader - Compact styling */
        [data-testid="stFileUploader"] {
            border: 2px dashed var(--info-blue) !important;
            border-radius: 8px !important;
            background: rgba(59, 130, 246, 0.03) !important;
            padding: 1rem !important;
            transition: all 0.2s ease !important;
            min-height: 60px !important;
            width: 100% !important;
        }
        
        [data-testid="stFileUploader"]:hover {
            border-color: var(--info-blue) !important;
            background: rgba(59, 130, 246, 0.05) !important;
            box-shadow: 0 2px 8px rgba(59, 130, 246, 0.1) !important;
        }
        
        /* Selected files preview */
        h3 {
            color: var(--primary-text) !important;
            font-weight: 600 !important;
            font-size: 1.2rem !important;
            margin-top: 1.5rem !important;
            margin-bottom: 1rem !important;
        }
        
        /* File uploader label - compact */
        [data-testid="stFileUploader"] label {
            color: var(--primary-text) !important;
            font-weight: 500 !important;
            font-size: 0.95rem !important;
            margin-bottom: 0.5rem !important;
        }
        
        /* File name container - simple, clean styling */
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] {
            background-color: var(--ui-panel) !important;
            border: 1px solid rgba(82, 82, 82, 0.3) !important;
            border-radius: 6px !important;
            padding: 0.5rem 0.75rem !important;
            margin: 0.5rem 0 !important;
            transition: all 0.15s ease !important;
            display: flex !important;
            align-items: center !important;
            justify-content: space-between !important;
            gap: 0.5rem !important;
            position: relative !important;
        }
        
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"]:hover {
            background-color: var(--ui-panel) !important;
            opacity: 0.9 !important;
            border-color: var(--info-blue) !important;
        }
        
        /* File name text - simple */
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] > *:first-child,
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] span:first-of-type,
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] div:first-of-type {
            color: var(--primary-text) !important;
            font-weight: 400 !important;
            font-size: 0.875rem !important;
            white-space: nowrap !important;
            overflow: hidden !important;
            text-overflow: ellipsis !important;
            flex: 1 !important;
            min-width: 0 !important;
        }
        
        /* File size - simple, aligned to right */
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileSize"],
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] + [data-testid="stFileUploaderFileSize"],
        [data-testid="stFileUploader"] span[data-testid="stFileUploaderFileSize"],
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] ~ span,
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] ~ div,
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] > [data-testid="stFileUploaderFileSize"],
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] span:last-child,
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] div:last-child:not(:first-child) {
            color: var(--secondary-text) !important;
            font-size: 0.8rem !important;
            font-weight: 400 !important;
            margin-left: auto !important;
            margin-right: 0.5rem !important;
            white-space: nowrap !important;
            flex-shrink: 0 !important;
        }
        
        /* File uploader browse button */
        [data-testid="stFileUploader"] button {
            background-color: transparent !important;
            border: none !important;
            color: var(--info-blue) !important;
            font-size: 1rem !important;
            padding: 0.25rem 0.5rem !important;
            cursor: pointer !important;
            transition: all 0.15s ease !important;
            border-radius: 4px !important;
        }
        
        [data-testid="stFileUploader"] button:hover {
            background-color: rgba(59, 130, 246, 0.1) !important;
            color: var(--info-blue) !important;
        }
        
        /* Success messages */
        .stSuccess {
            background-color: var(--verification-green);
            color: #ffffff;
        }
        
        /* Input field */
        .stTextInput > div > div > input {
            border-color: rgba(82, 82, 82, 0.3);
        }
        
        .stTextInput > div > div > input:focus {
            border-color: var(--info-blue);
            box-shadow: 0 0 0 2px rgba(59, 130, 246, 0.1);
        }
        
        /* Textarea styling - Cursor-style minimal design */
        .stTextArea > div > div > textarea {
            background-color: var(--ui-panel) !important;
            border: 1px solid rgba(82, 82, 82, 0.3) !important;
            border-radius: 6px !important;
            padding: 0.75rem 0.875rem !important;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif !important;
            font-size: 0.875rem !important;
            line-height: 1.6 !important;
            color: var(--primary-text) !important;
            box-shadow: none !important;
            transition: all 0.15s ease !important;
            resize: vertical !important;
            outline: none !important;
        }
        
        .stTextArea > div > div > textarea:focus {
            border-color: var(--info-blue) !important;
            box-shadow: none !important;
            outline: none !important;
            outline-width: 0 !important;
            outline-style: none !important;
            outline-color: transparent !important;
            background-color: var(--ui-panel) !important;
        }
        
        .stTextArea > div > div > textarea:focus-visible {
            outline: none !important;
            outline-width: 0 !important;
            outline-style: none !important;
            outline-color: transparent !important;
        }
        
        .stTextArea > div > div > textarea:hover {
            border-color: rgba(82, 82, 82, 0.5) !important;
        }
        
        /* Textarea container */
        .stTextArea > div {
            background-color: transparent !important;
        }
        
        /* Chat input */
        .stChatInputContainer {
            background-color: var(--ui-panel);
            border-top: 1px solid rgba(82, 82, 82, 0.3);
        }
        /* Citations/References - in main markdown */
        .stMarkdown strong {
            color: var(--primary-text);
        }
        
        /* References heading in assistant messages */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) 
        [data-testid="stChatMessageContent"] strong {
            color: var(--primary-text) !important;
        }
        
        /* Verification badges - pill-shaped indicators */
        .verification-badge {
            display: inline-block;
            padding: 3px 10px;
            margin-left: 6px;
            margin-right: 8px;
            border-radius: 12px;
            font-size: 11px;
            font-weight: 600;
            vertical-align: middle;
            line-height: 1.4;
            letter-spacing: 0.02em;
            text-transform: uppercase;
            min-width: 110px;
            width: auto;
            text-align: center;
            box-sizing: border-box;
            white-space: nowrap;
        }
        .verification-badge.verified {
            background-color: var(--verification-green);
            color: white;
        }
        .verification-badge.unverified {
            background-color: var(--warning-orange);
            color: white;
            position: relative;
            cursor: help;
        }
        .verification-badge.refuted {
            background-color: var(--user-error-red);
            color: white;
        }
        
        /* Force all text in assistant messages to use primary text color - override any other styles */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] p,
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] span:not(.verification-badge),
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] div:not(.verification-badge),
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] *:not(.verification-badge):not(.verification-badge *) {
            color: var(--primary-text) !important;
        }
        
        /* Tooltip for unverified/extrapolated badges */
        .verification-badge.unverified:hover::after {
            content: "This statement is part of the synthesized summary but could not be directly verified from the cited text. Please review the source for full context.";
            position: absolute;
            bottom: 100%;
            left: 50%;
            transform: translateX(-50%);
            margin-bottom: 8px;
            padding: 0.75rem 1rem;
            background-color: var(--primary-text);
            color: var(--background);
            font-size: 0.875rem;
            font-weight: 400;
            white-space: normal;
            width: 280px;
            border-radius: 6px;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15);
            z-index: 1000;
            pointer-events: none;
            text-transform: none;
            line-height: 1.5;
            text-align: left;
        }
        
        /* Tooltip arrow */
        .verification-badge.unverified:hover::before {
            content: "";
            position: absolute;
            bottom: 92%;
            left: 50%;
            transform: translateX(-50%);
            border: 6px solid transparent;
            border-top-color: var(--primary-text);
            z-index: 1001;
            pointer-events: none;
        }
        
        /* Badge in claim list - no left margin */
        .claim-item-badge {
            margin-left: 0 !important;
        }
        
        /* Links in citations */
        .stMarkdown a {
            color: var(--info-blue);
        }
        
        .stMarkdown a:hover {
            color: var(--info-blue);
            opacity: 0.8;
        }
        
        /* Reference list items - simple, clean styling */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) 
        [data-testid="stChatMessageContent"] ul {
            border-left: 2px solid rgba(82, 82, 82, 0.3);
            padding-left: 16px;
        }
        
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) 
        [data-testid="stChatMessageContent"] ol {
            padding-left: 16px;
        }
        
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) 
        [data-testid="stChatMessageContent"] li {
            color: var(--primary-text) !important;
            margin-bottom: 0.25rem;
        }
        
        /* Spinner */
        .stSpinner > div {
            border-color: var(--info-blue);
        }
        
        /* Dark mode specific overrides */
        @media (prefers-color-scheme: dark) {
            .stApp {
                background-color: var(--background);
            }
            
            .main .block-container {
                background-color: var(--background);
            }
            
            h1 {
                color: var(--primary-text) !important;
            }
            
            .stMarkdown p {
                color: var(--primary-text);
            }
            
            [data-testid="stSidebar"] {
                background-color: var(--ui-panel);
            }
            
            h2 {
                color: var(--primary-text) !important;
            }
            
            .stInfo {
                background-color: rgba(59, 130, 246, 0.15) !important;
                border-left-color: var(--info-blue) !important;
            }
            
            [data-testid="stFileUploader"] {
                background-color: var(--ui-panel);
                border-color: var(--secondary-text);
            }
            
            .drag-drop-overlay {
                background: rgba(59, 130, 246, 0.15);
                border-color: var(--info-blue);
            }
            
            .drag-drop-overlay-content {
                background: var(--ui-panel);
                border-color: var(--info-blue);
            }
            
            .drag-drop-overlay-content h2 {
                color: var(--info-blue);
            }
            
            .drag-drop-overlay-content p {
                color: var(--secondary-text);
            }
            
            /* Status text - dark mode */
            .status-update-text {
                color: var(--primary-text) !important;
            }
            
            /* Expander - dark mode */
            [data-testid="stExpander"] summary {
                color: var(--primary-text) !important;
            }
            
            [data-testid="stExpander"] summary:hover {
                background: rgba(255, 255, 255, 0.05) !important;
            }
            
            [data-testid="stExpander"] [data-testid="stExpanderContent"] {
                background: rgba(255, 255, 255, 0.03) !important;
                border-left: 2px solid rgba(170, 170, 170, 0.3) !important;
            }
            
            [data-testid="stExpander"][open] [data-testid="stExpanderContent"] {
                background: rgba(255, 255, 255, 0.05) !important;
                border-left: 2px solid rgba(170, 170, 170, 0.5) !important;
            }
            
            /* Expander in assistant messages */
            [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stExpander"] summary {
                color: var(--primary-text) !important;
            }
            
            [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stExpander"] summary:hover {
                background: rgba(255, 255, 255, 0.1) !important;
            }
            
            [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stExpander"] [data-testid="stExpanderContent"] {
                background: rgba(255, 255, 255, 0.05) !important;
                border-left: 2px solid rgba(170, 170, 170, 0.3) !important;
            }
            
            [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stExpander"][open] [data-testid="stExpanderContent"] {
                background: rgba(255, 255, 255, 0.08) !important;
                border-left: 2px solid rgba(170, 170, 170, 0.5) !important;
            }
            
            [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] {
                background-color: var(--ui-panel) !important;
                border-color: rgba(170, 170, 170, 0.3) !important;
            }
            
            [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"]:hover {
                background-color: var(--ui-panel) !important;
                opacity: 0.9 !important;
                border-color: var(--info-blue) !important;
            }
            
            [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] span,
            [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] div {
                color: var(--primary-text) !important;
            }
            
            /* Textarea dark mode styling - Cursor-style */
            .stTextArea > div > div > textarea {
                background-color: var(--ui-panel) !important;
                border: 1px solid rgba(170, 170, 170, 0.3) !important;
                color: var(--primary-text) !important;
                box-shadow: none !important;
                outline: none !important;
            }
            
            .stTextArea > div > div > textarea:focus {
                border-color: var(--info-blue) !important;
                box-shadow: none !important;
                outline: none !important;
                outline-width: 0 !important;
                outline-style: none !important;
                outline-color: transparent !important;
                background-color: var(--ui-panel) !important;
            }
            
            .stTextArea > div > div > textarea:focus-visible {
                outline: none !important;
                outline-width: 0 !important;
                outline-style: none !important;
                outline-color: transparent !important;
            }
            
            .stTextArea > div > div > textarea:hover {
                border-color: rgba(170, 170, 170, 0.5) !important;
            }
            
            .stChatInputContainer {
                background-color: var(--ui-panel);
                border-top: 1px solid rgba(170, 170, 170, 0.3);
            }
            
            /* Dark mode chat messages */
            [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="user"]) [data-testid="stChatMessageContent"] {
                background-color: var(--user-error-red);
            }
            
            [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatar"] [data-icon="assistant"]) [data-testid="stChatMessageContent"] {
                background-color: var(--ui-panel);
                color: var(--primary-text) !important;
            }
        }

        /* Hide the Streamlit header and footer */
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        
        /* Ensure sidebar toggle button remains visible and functional */
        [data-testid="stSidebarCollapseButton"],
        button[data-testid="stSidebarCollapseButton"],
        [aria-label*="sidebar"][aria-label*="toggle"],
        button[aria-label*="Open sidebar"],
        button[aria-label*="Close sidebar"] {
            visibility: hidden !important;
            display: none !important;
        }
        
        /* Ensure sidebar toggle button icon is visible */
        [data-testid="stSidebarCollapseButton"] svg,
        button[data-testid="stSidebarCollapseButton"] svg,
        [data-testid="stSidebarCollapseButton"]::before,
        [data-testid="stSidebarCollapseButton"]::after {
            display: block !important;
            visibility: visible !important;
        }
        
        /* Custom sidebar toggle icon styling */
        [data-testid="stSidebarCollapseButton"] .custom-sidebar-icon {
            width: 20px !important;
            height: 20px !important;
            fill: currentColor !important;
            stroke: currentColor !important;
            color: var(--primary-text) !important;
        }
        
        /* Hide text/aria-label and show icon */
        [data-testid="stSidebarCollapseButton"] .sr-only,
        [data-testid="stSidebarCollapseButton"] [class*="sr-only"],
        button[aria-label*="sidebar"] .sr-only {
            position: absolute !important;
            width: 1px !important;
            height: 1px !important;
            padding: 0 !important;
            margin: -1px !important;
            overflow: hidden !important;
            clip: rect(0, 0, 0, 0) !important;
            white-space: nowrap !important;
            border: 0 !important;
        }
        
        </style>
        """


def inject_custom_css() -> None:
    """Inject custom CSS for medical research UI with IBM brand colors.

    Optimized with caching to avoid re-injecting on every rerun.
    """
    css_content = get_css_content()
    # Use a key to ensure CSS is only injected once per session
    if "_css_injected" not in st.session_state:
        st.markdown(css_content, unsafe_allow_html=True)
        st.session_state["_css_injected"] = True
    else:
        # Still inject CSS but Streamlit will optimize it
        st.markdown(css_content, unsafe_allow_html=True)


def init_state() -> None:
    """Initialize Streamlit session state variables."""
    if "messages" not in st.session_state:
        st.session_state["messages"] = []
    if "ingested_docs" not in st.session_state:
        st.session_state["ingested_docs"] = []
    if "show_upload_ui" not in st.session_state:
        # Hide upload UI by default if documents exist, show if no documents
        has_docs = len(st.session_state.get("ingested_docs", [])) > 0
        st.session_state["show_upload_ui"] = not has_docs
    if "report_text" not in st.session_state:
        st.session_state["report_text"] = ""


def upload_section(ingestion: IngestionPipeline) -> Optional[List]:
    """Display file upload section in main area.

    Args:
        ingestion: IngestionPipeline instance for processing documents.

    Returns:
        List of uploaded files if any, None otherwise.
    """
    has_documents = len(st.session_state.get("ingested_docs", [])) > 0
    
    # Check if answer generation is in progress (status callback indicates generation)
    is_generating = "_status_callback" in st.session_state
    
    # If documents exist and upload UI is hidden, show toggle button
    if has_documents and not st.session_state.get("show_upload_ui", True):
        if st.button("Upload More Documents", type="secondary", use_container_width=False, disabled=is_generating):
            # Don't allow state change during generation
            if not is_generating:
                st.session_state["show_upload_ui"] = True
                # Use rerun for navigation state changes
                st.rerun()
        if is_generating:
            st.caption("⚠️ Answer generation in progress - please wait...")
        return None
    
    # Show upload UI
    st.header("Upload Documents")
    
    # Hide button if documents exist - positioned on the left
    if has_documents:
        if st.button("Hide Upload", type="secondary", use_container_width=False, disabled=is_generating):
            # Don't allow state change during generation
            if not is_generating:
                st.session_state["show_upload_ui"] = False
                # Use rerun for navigation state changes
                st.rerun()
        if is_generating:
            st.caption("⚠️ Answer generation in progress - please wait...")
        st.markdown("")  # Spacing after button
    
    # Upload area - full width
    uploaded_files = st.file_uploader(
        "Drag PDF files here or click to browse", 
        type=["pdf"], 
        accept_multiple_files=True,
        help="Upload one or more PDF documents for research. Files will be processed and indexed.",
        label_visibility="visible"
    )
    
    # Check if answer generation is in progress
    is_generating = "_status_callback" in st.session_state
    
    # Ingest button below uploader
    if uploaded_files:
        ingest_button = st.button(
            f"Ingest {len(uploaded_files)} file(s)", 
            type="primary",
            use_container_width=False,
            disabled=is_generating
        )
        if is_generating:
            st.caption("⚠️ Answer generation in progress - please wait before ingesting files...")
    else:
        ingest_button = False
    
    # Don't process during generation to avoid interruption
    if uploaded_files and ingest_button and not is_generating:
        with st.spinner(f"Processing {len(uploaded_files)} file(s)..."):
            for f in uploaded_files:
                doc_id = str(uuid.uuid4())
                with st.spinner(f"Uploading and ingesting {f.name}..."):
                    source_uri = ingestion.upload_to_cos(doc_id, f.name, f)
                    count = ingestion.ingest_pdf(doc_id, f.name, source_uri)
                # Store document info: (doc_id, filename, source_uri, count, title, author)
                # ingest_pdf returns (upserted_count, metadata_dict)
                if isinstance(count, tuple) and len(count) == 2:
                    chunk_count, metadata = count
                    title = metadata.get("title") if metadata else None
                    author = metadata.get("author") if metadata else None
                else:
                    chunk_count = count
                    title = None
                    author = None
                st.session_state["ingested_docs"].append((doc_id, f.name, source_uri, chunk_count, title, author))
        st.success(f"Successfully ingested {len(uploaded_files)} document(s)")
        # Hide upload UI after successful ingestion
        st.session_state["show_upload_ui"] = False
        # Rerun needed to show new documents and update UI
        st.rerun()
    
    return uploaded_files


def format_references_with_titles(sources: List[str]) -> str:
    """Format references with document titles and download links.

    Args:
        sources: List of source URIs to format.

    Returns:
        Formatted references string with document titles.
    """
    if not sources:
        return ""
    
    formatted_refs = []
    ingested_docs = st.session_state.get("ingested_docs", [])
    
    for source_uri in sources:
        # Find document info by source_uri
        doc_title = None
        doc_filename = None
        for doc_info in ingested_docs:
            if len(doc_info) >= 3 and doc_info[2] == source_uri:
                doc_filename = doc_info[1] if len(doc_info) >= 2 else None
                # Get title if available (position 4), fallback to filename
                if len(doc_info) >= 5:
                    doc_title = doc_info[4] or doc_filename
                else:
                    doc_title = doc_filename
                break
        
        # Create display name
        display_name = doc_title or doc_filename or source_uri.split("/")[-1] or "Document"
        
        # Create download link using Streamlit's download button approach
        # We'll use a data URI approach or create a download handler
        # For now, show the title with the source URI as a tooltip/link
        formatted_refs.append(f"- **{display_name}**")
    
    return "\n".join(formatted_refs)


def create_download_link(
    source_uri: str, ingestion: IngestionPipeline
) -> str:
    """Create a download link for a document from COS.

    Args:
        source_uri: S3 URI of the document.
        ingestion: IngestionPipeline instance for COS access.

    Returns:
        Download key for the document.
    """
    # Create a unique key for this download
    download_key = hashlib.md5(source_uri.encode()).hexdigest()[:16]
    
    # Store the source URI in session state for download handler
    if "download_cache" not in st.session_state:
        st.session_state["download_cache"] = {}
    st.session_state["download_cache"][download_key] = source_uri
    
    # Return download key that can be used with download button
    return download_key


def download_document_from_cos(
    download_key: str, ingestion: IngestionPipeline
) -> bytes:
    """Download a document from COS using the download key.

    Args:
        download_key: Download key stored in session state.
        ingestion: IngestionPipeline instance for COS access.

    Returns:
        Document content as bytes, empty bytes if not found.
    """
    if "download_cache" not in st.session_state:
        return b""
    
    source_uri = st.session_state["download_cache"].get(download_key)
    if not source_uri:
        return b""
    
    # Fetch from COS
    try:
        # Parse s3://bucket/key
        if not source_uri.startswith("s3://"):
            return b""
        
        _, rest = source_uri.split("s3://", 1)
        bucket, key = rest.split("/", 1)
        
        # Get object from COS
        obj = ingestion.cos.client.get_object(Bucket=bucket, Key=key)
        return obj["Body"].read()
    except Exception as e:
        st.error(f"Error downloading document: {e}")
        return b""


def _clean_report_flags() -> None:
    """Clear report flags for answers no longer in the report text.

    Removes report flags from session state for answers that have been
    removed from the report text by the user.
    """
    report_text = st.session_state.get("report_text", "")
    
    # Get all report flags from session state
    report_keys = [key for key in st.session_state.keys() if key.startswith("_report_")]
    
    # Check messages to find corresponding answers
    messages = st.session_state.get("messages", [])
    
    for key in report_keys:
        # Extract answer_hash from key (format: "_report_item_{hash}" or "_report_new_{hash}")
        if "_report_item_" in key:
            hash_str = key.replace("_report_item_", "")
        elif "_report_new_" in key:
            hash_str = key.replace("_report_new_", "")
        else:
            continue
        
        # Find the corresponding answer in messages
        answer_found = False
        for role, content in messages:
            if role == "assistant":
                # Try to extract answer from content (before References section)
                answer_only = content.split("**References:**")[0].strip() if "**References:**" in content else content.strip()
                
                # Get sources from content
                sources = []
                if "**References:**" in content:
                    import re
                    refs_text = content.split("**References:**", 1)[1]
                    source_uris = re.findall(r's3://[^\s\n]+', refs_text)
                    sources = source_uris
                
                # Calculate hash to match (use consistent string representation)
                answer_hash = hash((answer_only, str(sorted(sources))))
                hash_str_match = str(answer_hash)
                
                if hash_str_match == hash_str:
                    # Check if answer is still in report text
                    answer_snippet = answer_only[:100].strip() if len(answer_only) > 100 else answer_only.strip()
                    if answer_snippet and answer_snippet not in report_text:
                        # Answer no longer in report, clear the flag
                        st.session_state.pop(key, None)
                    answer_found = True
                    break
        
        # If flag exists but no corresponding message found, clear it
        if not answer_found:
            st.session_state.pop(key, None)


def add_to_report(
    answer: str, sources: List[str], query: str = ""
) -> None:
    """Add a verified answer to the report.

    Args:
        answer: The answer text to add.
        sources: List of source URIs for the answer.
        query: Original query (optional, not included in report).
    """
    if "report_text" not in st.session_state:
        st.session_state["report_text"] = ""
    
    # Format the entry (only answer and references, no query)
    # Add spacing only if report is not empty
    report_empty = not st.session_state["report_text"].strip()
    entry = "" if report_empty else "\n\n"
    # Strip leading/trailing whitespace from answer
    answer_clean = answer.strip()
    entry += f"{answer_clean}\n\n"
    
    if sources:
        entry += "**References:**\n"
        ingested_docs = st.session_state.get("ingested_docs", [])
        
        for i, source_uri in enumerate(sources, 1):
            # Find document title by source_uri
            doc_title = None
            for doc_info in ingested_docs:
                if len(doc_info) >= 3 and doc_info[2] == source_uri:
                    # Get title if available (position 4), fallback to filename
                    if len(doc_info) >= 5:
                        doc_title = doc_info[4]
                    if not doc_title and len(doc_info) >= 2:
                        doc_title = doc_info[1]  # Use filename as fallback
                    break
            
            # Use title if available, otherwise use source URI
            if doc_title:
                entry += f"{i}. **{doc_title}**\n"
            else:
                # Fallback to filename or source URI
                display_name = source_uri.split("/")[-1] if "/" in source_uri else source_uri
                entry += f"{i}. **{display_name}**\n"
    
    st.session_state["report_text"] += entry


def generate_bibliography(sources: List[str]) -> str:
    """Generate a formatted bibliography from source URIs.

    Args:
        sources: List of source URIs to include in bibliography.

    Returns:
        Formatted bibliography string in APA style.
    """
    if not sources:
        return ""
    
    bibliography = "\n\n## Bibliography\n\n"
    
    # Extract unique documents from sources
    unique_docs = {}
    ingested_docs = st.session_state.get("ingested_docs", [])
    
    for source in sources:
        # Find document metadata
        for doc_info in ingested_docs:
            if len(doc_info) >= 3 and doc_info[2] == source:
                title = doc_info[4] if len(doc_info) >= 5 and doc_info[4] else None
                author = doc_info[5] if len(doc_info) >= 6 and doc_info[5] else None
                filename = doc_info[1] if len(doc_info) >= 2 else None
                
                doc_key = source
                if doc_key not in unique_docs:
                    unique_docs[doc_key] = {
                        "title": title,
                        "author": author,
                        "filename": filename,
                        "source_uri": source
                    }
                break
    
    # Format in APA style
    for i, (doc_key, doc_info) in enumerate(unique_docs.items(), 1):
        # Use title if available, fallback to filename, then source URI
        title = doc_info.get("title")
        if not title:
            filename = doc_info.get("filename", "")
            if filename:
                title = filename.replace(".pdf", "").replace(".PDF", "")
            else:
                source_uri = doc_info.get("source_uri", "")
                title = source_uri.split("/")[-1] if "/" in source_uri else source_uri
        
        author = doc_info.get("author")
        if not author:
            author = "Unknown Author"
        
        # Format as APA citation: Author, A. A. (Year). Title. [Format]. Source
        # Since we don't have year, we'll use: Author, A. A. (n.d.). Title. [PDF document]
        # For multiple authors, format as: Author, A. A., & Author, B. B.
        # Check if author contains multiple authors (comma or "and")
        if "," in author or " and " in author.lower():
            # Multiple authors - format properly
            if " and " in author.lower():
                # Split by "and" and format
                authors = [a.strip() for a in re.split(r'\s+and\s+', author, flags=re.IGNORECASE)]
                if len(authors) == 2:
                    formatted_author = f"{authors[0]}, & {authors[1]}"
                else:
                    # More than 2 authors - last one gets "&"
                    formatted_author = ", ".join(authors[:-1]) + ", & " + authors[-1]
            else:
                # Comma-separated - assume already formatted
                formatted_author = author
        else:
            # Single author
            formatted_author = author
        
        # APA format: Author, A. A. (n.d.). Title. [PDF document].
        bibliography += f"{i}. {formatted_author} (n.d.). {title}. [PDF document].\n"
    
    return bibliography


def export_report(format: str = "docx") -> bytes:
    """Export the research report in the specified format.

    Args:
        format: Export format, either "docx" or "md".

    Returns:
        Report content as bytes, empty bytes if report is empty.
    """
    report_text = st.session_state.get("report_text", "")
    if not report_text:
        return b""
    
    # Extract sources from report text
    source_pattern = r's3://[^\s\n]+'
    all_sources = re.findall(source_pattern, report_text)
    unique_sources = list(dict.fromkeys(all_sources))
    
    if format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            
            # Title
            doc.add_heading('Synthesis Studio', 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph("")  # Blank line
            
            # Parse report text and convert to DOCX
            lines = report_text.split('\n')
            current_para = []
            in_refs_section = False
            
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    if current_para:
                        para = doc.add_paragraph(" ".join(current_para))
                        current_para = []
                elif line_stripped.startswith("**References:**"):
                    if current_para:
                        para = doc.add_paragraph(" ".join(current_para))
                        current_para = []
                    doc.add_heading("References", level=2)
                    in_refs_section = True
                elif in_refs_section and re.match(r'^\d+\.', line_stripped):
                    # Reference list item
                    ref_text = line_stripped
                    doc.add_paragraph(ref_text, style='List Number')
                else:
                    # Regular paragraph text
                    if line_stripped.startswith("**") and line_stripped.endswith("**"):
                        # Bold text
                        current_para.append(line_stripped.replace("**", ""))
                    else:
                        current_para.append(line_stripped)
            
            if current_para:
                doc.add_paragraph(" ".join(current_para))
            
            # Add bibliography
            bibliography = generate_bibliography(unique_sources)
            if bibliography:
                doc.add_page_break()
                for line in bibliography.split('\n'):
                    if line.strip():
                        if line.startswith("##"):
                            doc.add_heading(line.replace("##", "").strip(), level=2)
                        elif re.match(r'^\d+\.', line.strip()):
                            doc.add_paragraph(line.strip(), style='List Number')
                        else:
                            doc.add_paragraph(line.strip())
            
            # Save to bytes
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
        except ImportError:
            return b""
    
    elif format == "md":
        md_content = report_text
        bibliography = generate_bibliography(unique_sources)
        if bibliography:
            md_content += bibliography
        
        return md_content.encode('utf-8')
    
    return b""


def report_page() -> None:
    """Display and manage the research report page.

    Provides UI for editing, exporting, and managing the synthesis report.
    """
    st.header("Synthesis Studio")
    st.caption("Build your final deliverable from verified insights")
    
    # Export buttons at the top
    try:
        from docx import Document
        DOCX_AVAILABLE = True
    except ImportError:
        DOCX_AVAILABLE = False
    
    col1, col2 = st.columns([1, 4])
    
    with col1:
        if DOCX_AVAILABLE:
            docx_data = export_report("docx")
            st.download_button(
                label="Export as DOCX",
                data=docx_data if docx_data else b"",
                file_name=f"research_report_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                disabled=not st.session_state.get("report_text", ""),
                use_container_width=True
            )
        else:
            st.info("⚠️ python-docx not available. Install it to export as DOCX.")
    
    with col2:
        md_data = export_report("md")
        st.download_button(
            label="Export as Markdown",
            data=md_data if md_data else b"",
            file_name=f"research_report_{datetime.now().strftime('%Y%m%d')}.md",
            mime="text/markdown",
            disabled=not st.session_state.get("report_text", ""),
            use_container_width=True
        )
    
    st.divider()
    
    # Report text area (editable)
    if not st.session_state.get("report_text", ""):
        st.info("Your report is empty. Add verified findings from the MedCortex Analyst page to build your report.")
        st.session_state["report_text"] = ""
    
    # Editable text area
    report_content = st.text_area(
        "Synthesis Content",
        value=st.session_state.get("report_text", ""),
        height=500,
        help="Edit your synthesis here. Content added via 'Add to Studio' buttons will appear here.",
        label_visibility="collapsed"
    )
    
    # Update session state if user edits
    if report_content != st.session_state.get("report_text", ""):
        st.session_state["report_text"] = report_content
        # Clear report flags for answers that are no longer in the report
        _clean_report_flags()
    
    # Clear button
    if st.button("Clear Report", use_container_width=True):
        st.session_state["report_text"] = ""
        # Clear all report flags when clearing report
        _clean_report_flags()
        # Rerun needed to clear the text area
        st.rerun()


def sidebar_documents() -> None:
    """Display compact sidebar showing ingested documents.

    Shows list of ingested documents with their chunk counts.
    """
    st.sidebar.markdown("### Ingested Documents")
    
    ingested_docs = st.session_state.get("ingested_docs", [])
    if ingested_docs:
        st.sidebar.caption(f"{len(ingested_docs)} document(s)")
        # Updated to handle new format: (doc_id, filename, source_uri, count, title, author)
        for doc_info in ingested_docs:
            if len(doc_info) >= 4:
                doc_id, filename, uri, count = doc_info[0], doc_info[1], doc_info[2], doc_info[3]
                
                # Prefer title over filename, with fallback to filename
                title = doc_info[4] if len(doc_info) >= 5 and doc_info[4] else None
                display_name = title or filename
                
                # Smart truncation: truncate at word boundaries if possible
                max_length = 40  # Maximum characters to display
                if len(display_name) > max_length:
                    # Try to truncate at word boundary
                    truncated = display_name[:max_length]
                    last_space = truncated.rfind(' ')
                    if last_space > max_length * 0.7:  # If space is reasonably close to max
                        display_name = truncated[:last_space] + "..."
                    else:
                        # Just truncate at max_length
                        display_name = truncated + "..."
                
                st.sidebar.markdown(f"**{display_name}**")
                st.sidebar.caption(f"{count} chunks")
    else:
        st.sidebar.info("Upload documents to get started")


def display_agent_trajectory(
    query: str, trajectory: List[dict]
) -> None:
    """Display the agent's step-by-step reasoning process.

    Args:
        query: Original user query.
        trajectory: List of trajectory steps showing reasoning process.
    """
    with st.expander("Analysis Breakdown", expanded=False):
        st.caption(f"Showing step-by-step reasoning for: \"{query[:30]}...{query[-30:] if len(query) > 60 else query}\"")
        
        with st.container():
            for step_info in trajectory:
                step_type = step_info.get("type", "")
                title = step_info.get("title", "")
                content = step_info.get("content", "")
                details = step_info.get("details", "")
                
                # Style based on step type
                if step_type == "planning":
                    with st.expander(f"{title}", expanded=False):
                        st.info(content)
                        if details:
                            st.caption(details)
                elif step_type == "decomposition":
                    with st.expander(f"{title}", expanded=False):
                        st.markdown(content)
                        if details:
                            st.caption(details)
                elif step_type == "retrieval":
                    with st.expander(f"{title}", expanded=False):
                        st.markdown(f"**Query:** {content}")
                        if details:
                            st.caption(details)
                elif step_type == "intermediate_answer":
                    with st.expander(f"{title}", expanded=False):
                        st.markdown(content)
                        if details:
                            st.caption(details)
                        # Show full answer in expander
                        full_answer = step_info.get("full_answer", "")
                        if full_answer and len(full_answer) > len(content):
                            with st.expander("View full intermediate answer", expanded=False):
                                st.markdown(full_answer)
                        # Show sources if available
                        sources = step_info.get("sources", [])
                        if sources:
                            st.caption(f"Sources: {len(sources)}")
                elif step_type == "synthesis":
                    with st.expander(f"{title}", expanded=False):
                        st.markdown(content)
                        if details:
                            st.caption(details)
                elif step_type == "verification":
                    with st.expander(f"{title}", expanded=False):
                        st.markdown(content)
                        if details:
                            st.caption(details)
                elif step_type == "verification_result":
                    with st.expander(f"{title}", expanded=False):
                        st.markdown(content)
                        if details:
                            st.caption(details)
                        # Show verification details if available
                        verification_data = step_info.get("verification_results", [])
                        if verification_data:
                            supports = [r for r in verification_data if r.get("status") == "Supports"]
                            refutes = [r for r in verification_data if r.get("status") == "Refutes"]
                            not_mentioned = [r for r in verification_data if r.get("status") == "Not Mentioned"]
                            
                            if supports:
                                st.success(f"{len(supports)} claim(s) verified against sources")
                            if refutes:
                                st.error(f"{len(refutes)} claim(s) contradicted by sources")
                            if not_mentioned:
                                st.warning(f"{len(not_mentioned)} claim(s) extrapolated from sources")
                elif step_type == "final_answer":
                    st.success(f"{title}")
                    if details:
                        st.caption(details)


def display_answer_with_verification(
    answer_text: str, verification_results: List[dict]
) -> None:
    """Display answer with verification status indicators.

    Args:
        answer_text: The answer text to display.
        verification_results: List of verification results with claim statuses.
    """
    if not verification_results:
        st.markdown(answer_text)
        return
    
    # Create a mapping of claims to status
    claim_to_status = {r.get("claim", ""): r.get("status", "Not Mentioned") for r in verification_results}
    
    # Track which claims have been matched to avoid duplicate badges
    matched_claims = set()
    
    # Clean answer text (remove citations section for processing)
    answer_only = answer_text.split("**References:**")[0].strip() if "**References:**" in answer_text else answer_text
    
    # Split answer into sentences and add verification markers
    sentences = re.split(r'([.;]\s+|\n+)', answer_only)
    
    displayed_text = ""
    for sentence in sentences:
        sentence_clean = sentence.strip()
        if not sentence_clean or len(sentence_clean) < 10:
            displayed_text += sentence
            continue
        
        # Check if any claim matches this sentence (fuzzy match)
        # Only match claims that haven't been matched yet
        status = None
        best_match = None
        best_match_score = 0
        
        for claim, claim_status in claim_to_status.items():
            if claim and len(claim) > 15 and claim not in matched_claims:
                # Check if claim text appears in sentence (first 50 chars for matching)
                claim_preview = claim[:50].lower().strip()
                sentence_preview = sentence_clean[:100].lower().strip()
                
                match_score = 0
                # Check for substring match (exact match is best)
                if claim_preview in sentence_preview:
                    match_score = 100  # High score for substring match
                    status = claim_status
                    best_match = claim
                    best_match_score = match_score
                    break  # Exact match, stop searching
                # Check word overlap for longer claims
                claim_words = set(w for w in claim_preview.split() if len(w) > 3)
                sentence_words = set(w for w in sentence_preview.split() if len(w) > 3)
                if claim_words and len(claim_words & sentence_words) >= 2:
                    # Calculate overlap score (more overlap = higher score)
                    overlap_score = len(claim_words & sentence_words) / len(claim_words) if claim_words else 0
                    match_score = overlap_score * 50  # Medium score for word overlap
                    if match_score > best_match_score:
                        status = claim_status
                        best_match = claim
                        best_match_score = match_score
        
        # Only add badge if we found a match and it hasn't been used yet
        if best_match and best_match not in matched_claims:
            matched_claims.add(best_match)
            if status == "Supports":
                # Add verified badge
                displayed_text += f'{sentence}<span class="verification-badge verified">Verified</span>'
            elif status == "Refutes":
                # Add refuted badge
                displayed_text += f'{sentence}<span class="verification-badge refuted">Refuted</span>'
            elif status == "Not Mentioned":
                # Add unverified badge
                displayed_text += f'{sentence}<span class="verification-badge unverified">Extrapolated</span>'
            else:
                displayed_text += sentence
        else:
            displayed_text += sentence
    
    st.markdown(displayed_text, unsafe_allow_html=True)
    
    # Show verification summary in expander
    with st.expander("Verification Details", expanded=False):
        supports = [r for r in verification_results if r.get("status") == "Supports"]
        refutes = [r for r in verification_results if r.get("status") == "Refutes"]
        not_mentioned = [r for r in verification_results if r.get("status") == "Not Mentioned"]
        
        # Create styled summary badges
        summary_html = "<div style='display: flex; flex-wrap: wrap; gap: 8px; margin-bottom: 16px;'>"
        
        if supports:
            summary_html += f'<span class="verification-badge verified">{len(supports)} Verified</span>'
        if refutes:
            summary_html += f'<span class="verification-badge refuted">{len(refutes)} Refuted</span>'
        if not_mentioned:
            summary_html += f'<span class="verification-badge unverified">{len(not_mentioned)} Extrapolated</span>'
        
        summary_html += "</div>"
        st.markdown(summary_html, unsafe_allow_html=True)
        
        if supports:
            st.info(f"**{len(supports)} claim(s) verified** against source documents")
        if refutes:
            st.error(f"**{len(refutes)} claim(s) contradicted** by source documents")
        if not_mentioned:
            st.warning(f"**{len(not_mentioned)} claim(s) extrapolated** from source documents")
        
        # Show individual claim details
        if verification_results:
            st.divider()
            st.caption("Individual Claims:")
            for i, result in enumerate(verification_results, 1):
                claim = result.get("claim", "")  # Show full claim, no truncation
                status = result.get("status", "Not Mentioned")
                
                if status == "Supports":
                    badge_class = "verification-badge verified claim-item-badge"
                elif status == "Refutes":
                    badge_class = "verification-badge refuted claim-item-badge"
                else:
                    badge_class = "verification-badge unverified claim-item-badge"
                
                badge_html = f'<span class="{badge_class}">{"Verified" if status == "Supports" else "Refuted" if status == "Refutes" else "Extrapolated"}</span>'
                
                # Use flexbox layout to ensure wrapped text aligns with "Claim X:" label
                claim_html = f"""
                <div style="display: flex; align-items: flex-start; margin-bottom: 0.75rem; gap: 8px;">
                    <div style="flex-shrink: 0;">
                        {badge_html}
                    </div>
                    <div style="flex: 1; min-width: 0;">
                        <strong>Claim {i}:</strong> {claim}
                    </div>
                </div>
                """
                st.markdown(claim_html, unsafe_allow_html=True)


def chat_ui(
    query_pipeline: QueryPipeline,
    ingestion: Optional[IngestionPipeline] = None,
) -> None:
    """Main chat interface for querying documents.

    Args:
        query_pipeline: QueryPipeline instance for processing queries.
        ingestion: Optional IngestionPipeline for download functionality.
    """
    st.header("MedCortex Analyst")
    
    if not st.session_state["ingested_docs"]:
        st.info("Upload documents above to start asking questions about your research materials.")
        # Show placeholder chat interface even without documents
        return
    
    # Get document IDs for current session to filter search results
    # Updated to handle new format: (doc_id, filename, source_uri, count, title, author)
    session_doc_ids = []
    for doc_info in st.session_state["ingested_docs"]:
        if len(doc_info) >= 4:
            session_doc_ids.append(doc_info[0])  # doc_id is first element
    
    # Display chat history with verification status and trajectory
    # Use cached rendering to avoid unnecessary recomputation
    messages = st.session_state.get("messages", [])
    for idx, (role, content) in enumerate(messages):
        with st.chat_message(role):
            if role == "assistant":
                # Extract answer and sources from content first
                answer_only = content.split("**References:**")[0].strip() if "**References:**" in content else content
                
                # Extract sources from content
                sources = []
                if "**References:**" in content:
                    refs_text = content.split("**References:**", 1)[1].strip()
                    # Extract sources from reference list (handles both bullet points and numbered lists)
                    import re
                    # Match s3:// URIs, handling both "- s3://..." and "1. s3://..." formats
                    sources = re.findall(r's3://[^\s\n]+', refs_text)
                    sources = [s.rstrip('-').strip().lstrip('0123456789. ').strip() for s in sources]
                    sources = [s for s in sources if s.startswith('s3://')]  # Filter to only valid URIs
                
                # Check for trajectory first
                trajectory_shown = False
                if "agent_trajectory" in st.session_state and st.session_state["agent_trajectory"]:
                    # Find trajectory for this message (match by answer content)
                    for traj_entry in st.session_state["agent_trajectory"]:
                        # Try to match by checking if the answer is in the content
                        if traj_entry.get("answer") and traj_entry["answer"] in content:
                            trajectory_data = traj_entry.get("trajectory")
                            if trajectory_data:
                                display_agent_trajectory(traj_entry.get("query", ""), trajectory_data)
                                trajectory_shown = True
                                break
                
                # Check if verification results are available
                verification_results = []
                if "verification_results" in st.session_state:
                    verification_data = st.session_state["verification_results"]
                    for verif in verification_data:
                        if verif.get("answer") == answer_only:
                            verification_results = verif.get("verification", [])
                            # Use sources from verification if available, otherwise keep extracted sources
                            sources = verif.get("sources", sources)
                            break
                
                if verification_results:
                    display_answer_with_verification(answer_only, verification_results)
                    # Show references separately if they exist
                    if "**References:**" in content:
                        # Extract and format references with titles
                        refs_text = content.split("**References:**", 1)[1].strip()
                        import re
                        # Match s3:// URIs
                        source_uris = re.findall(r's3://[^\s\n]+', refs_text)
                        source_uris = [s.rstrip('-').strip().lstrip('0123456789. ').strip() for s in source_uris]
                        source_uris = [s for s in source_uris if s.startswith('s3://')]
                        
                        if source_uris:
                            st.markdown("\n\n**References:**")
                            ingested_docs = st.session_state.get("ingested_docs", [])
                            for source_uri in source_uris:
                                # Find document info by source_uri
                                doc_title = None
                                for doc_info in ingested_docs:
                                    if len(doc_info) >= 3 and doc_info[2] == source_uri:
                                        # Get title if available (position 4)
                                        if len(doc_info) >= 5:
                                            doc_title = doc_info[4]
                                        break
                                
                                # Create display text with title only (authors only used for bibliography)
                                if doc_title:
                                    display_text = f"- **{doc_title}**"
                                else:
                                    # Fallback to filename if no title
                                    for doc_info in ingested_docs:
                                        if len(doc_info) >= 3 and doc_info[2] == source_uri:
                                            doc_filename = doc_info[1] if len(doc_info) >= 2 else None
                                            display_text = f"- **{doc_filename or 'Document'}**"
                                            break
                                    else:
                                        display_text = f"- **{source_uri.split('/')[-1] or 'Document'}**"
                                
                                st.markdown(display_text)
                        else:
                            # Fallback to original display
                            st.markdown(f"\n\n**References:**\n{refs_text}")
                else:
                    # No verification found, just display the content
                    if not trajectory_shown:
                        st.markdown(content)
                
                # Add "Add to Report" button inside chat message context
                # Use sources extracted above
                sources_final = sources
                
                # Extract query from previous user message
                query_text = ""
                if idx > 0 and messages[idx - 1][0] == "user":
                    query_text = messages[idx - 1][1]
                
                # Place button inside the chat message, after content
                button_key = f"add_report_{idx}"
                # Use a stable key based on answer content hash
                answer_hash = hash((answer_only, str(sorted(sources_final))))
                report_key = f"_report_item_{answer_hash}"
                
                # Check flag first (more reliable than text search)
                is_flagged = st.session_state.get(report_key, False)
                
                # Check if answer is already in the report text
                report_text = st.session_state.get("report_text", "")
                # Use first 100 chars of answer to check if it's in report (more reliable than full text)
                answer_snippet = answer_only[:100].strip() if len(answer_only) > 100 else answer_only.strip()
                is_in_text = answer_snippet in report_text if answer_snippet else False
                
                # If flag is set but answer is not in text, clear the flag (user deleted it)
                if is_flagged and not is_in_text:
                    st.session_state.pop(report_key, None)
                    is_flagged = False
                
                is_already_added = is_flagged or is_in_text
                
                if is_already_added:
                    # Show disabled button if already added
                    st.button("Already in Synthesis Studio", key=f"report_status_{idx}", disabled=True, use_container_width=False)
                else:
                    # Show active button only if not already added
                    button_clicked = st.button("Add to Synthesis Studio", key=button_key, use_container_width=False)
                    
                    if button_clicked:
                        # Set flag first to prevent duplicate clicks
                        st.session_state[report_key] = True
                        # Add to report immediately
                        add_to_report(answer_only, sources_final, query_text)
                        # Rerun to update button state
                        st.rerun()
            else:
                st.markdown(content)

    # Chat input
    user_input = st.chat_input("Ask a research question about your documents…")
    
    if user_input:
        st.session_state["messages"].append(("user", user_input))
        with st.chat_message("user"):
            st.markdown(user_input)

        with st.chat_message("assistant"):
            # Check if this query used the orchestrator (has trajectory)
            trajectory_data = None
            
            # Create dynamic status display using empty container
            status_container = st.empty()
            
            def update_status(status_text: str):
                """Update the status text dynamically"""
                # Use status container to update text with pulsing animation (no background, no emoji)
                with status_container.container():
                    st.markdown(f'<p class="status-update-text">{status_text}</p>', unsafe_allow_html=True)
            
            # Store status callback in session state so orchestrator/pipeline can use it
            st.session_state["_status_callback"] = update_status
            
            # Start with initial status
            update_status("Searching documents and generating answer...")
            
            # Filter search to only use documents from current session
            answer, sources = query_pipeline.answer(user_input, allowed_doc_ids=session_doc_ids)
            
            # Clear status and callback after processing
            status_container.empty()
            if "_status_callback" in st.session_state:
                del st.session_state["_status_callback"]
                
                # Get trajectory if available (check after answer is generated)
                if "agent_trajectory" in st.session_state and st.session_state["agent_trajectory"]:
                    # Get the latest trajectory
                    latest_traj = st.session_state["agent_trajectory"][-1]
                    if latest_traj.get("query") == user_input:
                        trajectory_data = latest_traj.get("trajectory")
            
            # Show trajectory if available (display before answer)
            if trajectory_data:
                display_agent_trajectory(user_input, trajectory_data)
            
            # Get verification results if available
            verification_results = []
            if "verification_results" in st.session_state and st.session_state["verification_results"]:
                # Get the latest verification result
                latest_verif = st.session_state["verification_results"][-1]
                if latest_verif.get("answer") == answer:
                    verification_results = latest_verif.get("verification", [])
            
            # Display answer with verification
            display_answer_with_verification(answer, verification_results)
            
            # Append citations separately with titles only
            if sources:
                # Display references with titles only (authors only used for bibliography)
                st.markdown("\n\n**References:**")
                refs_list = []
                ingested_docs = st.session_state.get("ingested_docs", [])
                
                for source_uri in sources:
                    # Find document info by source_uri
                    doc_title = None
                    for doc_info in ingested_docs:
                        if len(doc_info) >= 3 and doc_info[2] == source_uri:
                            # Get title if available (position 4), fallback to filename
                            if len(doc_info) >= 5:
                                doc_title = doc_info[4]
                            break
                    
                    # Create display text with title only
                    if doc_title:
                        display_text = f"- **{doc_title}**"
                    else:
                        # Fallback to filename if no title
                        for doc_info in ingested_docs:
                            if len(doc_info) >= 3 and doc_info[2] == source_uri:
                                doc_filename = doc_info[1] if len(doc_info) >= 2 else None
                                display_text = f"- **{doc_filename or 'Document'}**"
                                break
                        else:
                            display_text = f"- **{source_uri.split('/')[-1] or 'Document'}**"
                    
                    st.markdown(display_text)
                    refs_list.append(display_text)
                
                citations = "\n\n**References:**\n" + "\n".join(refs_list)
                full_response = answer + citations
            else:
                full_response = answer
            
            # Add "Add to Report" button inside chat message context
            # Use a stable key based on answer content hash
            answer_hash = hash((answer, str(sorted(sources))))
            report_key = f"_report_new_{answer_hash}"
            
            # Check flag first (more reliable than text search)
            is_flagged = st.session_state.get(report_key, False)
            
            # Check if answer is already in the report text
            report_text = st.session_state.get("report_text", "")
            # Use first 100 chars of answer to check if it's in report (more reliable than full text)
            answer_snippet = answer[:100].strip() if len(answer) > 100 else answer.strip()
            is_in_text = answer_snippet in report_text if answer_snippet else False
            
            # If flag is set but answer is not in text, clear the flag (user deleted it)
            if is_flagged and not is_in_text:
                st.session_state.pop(report_key, None)
                is_flagged = False
            
            is_already_added = is_flagged or is_in_text
            
            if is_already_added:
                # Show disabled button if already added
                st.button("Already in Synthesis Studio", key="report_status_new", disabled=True, use_container_width=False)
            else:
                # Show active button only if not already added
                button_clicked = st.button("Add to Synthesis Studio", key="add_report_new", use_container_width=False)
                
                if button_clicked:
                    # Set flag first to prevent duplicate clicks
                    st.session_state[report_key] = True
                    # Add to report immediately
                    add_to_report(answer, sources, user_input)
                    # Rerun to update button state
                    st.rerun()
            
            st.session_state["messages"].append(("assistant", full_response))


def research_assistant_page(
    ingestion: IngestionPipeline, query_pipeline: QueryPipeline
) -> None:
    """Display the Research Assistant page.

    Args:
        ingestion: IngestionPipeline for document processing.
        query_pipeline: QueryPipeline for query processing.
    """
    st.title("MedCortex")
    st.caption("AI Research Analyst • Verifiable Synthesis • Powered by watsonx.ai")
    
    # Disclaimer under title
    st.markdown(
        '<p style="font-size: 0.85rem; color: var(--secondary-text); opacity: 0.6; margin-top: -0.8rem; margin-bottom: 1rem;">'
        'While MedCortex strives for accuracy through verification and source attribution, '
        'AI-generated content may contain errors. Please review all information and verify '
        'critical findings against original sources.'
        '</p>',
        unsafe_allow_html=True
    )
    
    # Upload section in main area (can be hidden)
    upload_section(ingestion)
    
    # Chat interface (pass ingestion for download functionality)
    chat_ui(query_pipeline, ingestion)


def research_report_page() -> None:
    """Display the Research Report page.

    Shows the Synthesis Studio for building the final deliverable.
    """
    st.title("MedCortex")
    st.caption("Synthesis Studio • Your workspace for curating evidence-based analysis")
    
    report_page()


def about_page() -> None:
    """Display the About page.

    Shows information about MedCortex, its mission, and features.
    """
    st.title("About MedCortex")
    
    st.markdown("""
    MedCortex was built to solve the single biggest challenge in modern medical research: **the synthesis headache**. 
    Researchers, clinicians, and academics are overwhelmed by the sheer volume of literature. Finding information is hard, 
    but synthesizing it—connecting data, comparing findings, and building a trusted evidence base—is a slow, manual, 
    and high-anxiety process.
    
    Standard AI tools promise speed but lack the required rigor. A "black box" answer is useless in a field that runs on 
    evidence. Hallucinations aren't just errors; they're a risk.
    """)
    
    st.header("Our Solution: An Analyst, Not a Search Bar")
    
    st.markdown("""
    MedCortex is **not a search engine**. It's an **AI research analyst**.
    
    Based on a state-of-the-art agentic framework and powered by IBM watsonx.ai, MedCortex performs the real work of 
    synthesis. When you ask a complex question, it doesn't just "find" an answer; it creates a plan. You can see this in 
    the **"Analysis Breakdown,"** where the agent performs Query Analysis and Query Decomposition, executing a multi-step 
    process just as an expert human researcher would.
    
    It intelligently routes your query to the right tool, performing advanced **Hybrid Search** for textual concepts and 
    separate, structured-data analysis for information locked in tables (**TableRAG**).
    """)
    
    
    st.header("The Core of MedCortex: Trust Through Verification")
    
    st.markdown("""
    Our **"Analysis Breakdown"** creates transparency. Our **"Verification Engine"** builds trust. MedCortex is a **"glass box."**
    
    After generating an answer, the platform performs a crucial third step: it fact-checks every single claim against its 
    source documents. Findings are clearly marked in the chat with a **"VERIFIED"** tag. If a claim is an AI-generated summary 
    that cannot be directly supported by the text, it is explicitly labeled as **"REFUTED,"** giving you full control and 
    transparency.
    """)
    
    
    st.header("From Analysis to Deliverable")
    
    st.markdown("""
    MedCortex is designed to fit your professional workflow, from initial **"Objectives"** to final **"Deliverables"**. 
    As you gather verified insights from the Analyst Chat, you can add them to your **"Synthesis Studio."** This curated 
    workspace is where your analysis becomes a formatted research report, complete with citations. You can then export your 
    work as a `.docx` or Markdown file, turning days of manual writing into minutes of curation.
    """)
    
    st.markdown("""
    ### This is MedCortex: Your AI Research Analyst for verifiable, end-to-end synthesis.
    """)
    
    st.title("About the Creator")
    
    st.markdown("""
    MedCortex is built by **Rohan Ramakrishnan**, a student at the University of Southern California (USC) pursuing a unique, interdisciplinary 
    blend of Computer Science and Business Administration.
    
    This project was born from a deep passion for improving human well-being and a great appreciation for the power of new 
    research in healthcare. The goal was to build a tool that directly addresses the "synthesis headache" by moving beyond 
    simple search to provide a complete, end-to-end solution for researchers.
    
    This approach is informed by his professional experience as a Founding Engineer at an artist super-fan platform and as 
    a Software Engineer for HackSC. In these roles, he developed a deep knowledge of building end-to-end systems that solve 
    complex, human-facing problems.
    
    MedCortex is the direct application of that end-to-end philosophy. It is designed as a complete, professional workspace 
    that aligns with the "Objective-driven" and "Phase"-based process of a real research plan. It guides a user from the 
    rigorous "Analysis Breakdown" of complex queries all the way to a final, verified "Deliverable" in the "Synthesis Studio".
    
    When not in front of a keyboard, he is an avid Formula 1 enthusiast, appreciating the blend of high-performance 
    engineering, data analytics, and mission-driven teamwork that defines the sport.
    """)


def research_about_page() -> None:
    """Display the About page.

    Shows information about MedCortex, its mission, and features.
    """
    st.title("MedCortex")
    st.caption("Our Mission & Method • How the 'Analysis Breakdown' delivers verified deliverables")
    
    about_page()


@st.cache_resource
def get_settings() -> Settings:
    """Load and cache application settings.

    Returns:
        Settings instance loaded from environment variables.
    """
    load_dotenv()
    return Settings.from_env()


@st.cache_resource
def get_pipelines(
    settings: Settings,
) -> Tuple[IngestionPipeline, QueryPipeline]:
    """Initialize and cache pipelines.

    Args:
        settings: Application settings.

    Returns:
        Tuple of (IngestionPipeline, QueryPipeline) instances.
    """
    return IngestionPipeline(settings), QueryPipeline(settings)


def main() -> None:
    """Main entry point for the Streamlit application.

    Sets up the page configuration, initializes pipelines, and handles
    navigation between different pages (Assistant and Report).
    """
    # Set page config for MedCortex branding
    st.set_page_config(
        page_title="MedCortex | Analysis Workspace",
        page_icon=None,
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Load settings (cached)
    settings = get_settings()
    init_state()
    
    # Inject custom CSS for medical research UI (optimized)
    inject_custom_css()

    # Get pipelines (cached)
    ingestion, query_pipeline = get_pipelines(settings)

    # Initialize navigation state
    if "current_page" not in st.session_state:
        st.session_state["current_page"] = "assistant"
    
    # Navigation in sidebar - separate pages
    st.sidebar.markdown("### Navigation")
    
    # Get current page from query params or session state
    try:
        query_params = st.query_params
        if "page" in query_params:
            current_page = query_params["page"]
            st.session_state["current_page"] = current_page
        else:
            current_page = st.session_state.get("current_page", "assistant")
    except:
        # Fallback if query_params not available
        current_page = st.session_state.get("current_page", "assistant")
    
    # Check if answer generation is in progress
    is_generating = "_status_callback" in st.session_state
    
    # Navigation links - optimized to reduce reruns
    # Disable navigation during generation to prevent interruption
    nav_assistant = st.sidebar.button("MedCortex Analyst", key="nav_assistant", use_container_width=True, type="secondary", disabled=is_generating)
    nav_report = st.sidebar.button("Synthesis Studio", key="nav_report", use_container_width=True, type="secondary", disabled=is_generating)
    nav_about = st.sidebar.button("Platform Overview", key="nav_about", use_container_width=True, type="secondary", disabled=is_generating)
    
    if is_generating:
        st.sidebar.caption("⚠️ Answer generation in progress...")
    
    # Handle navigation - only rerun if page actually changes and not during generation
    if nav_assistant and not is_generating:
        if st.session_state.get("current_page") != "assistant":
            st.session_state["current_page"] = "assistant"
            try:
                st.query_params.page = "assistant"
            except:
                pass
            st.rerun()
    
    if nav_report and not is_generating:
        if st.session_state.get("current_page") != "report":
            st.session_state["current_page"] = "report"
            try:
                st.query_params.page = "report"
            except:
                pass
            st.rerun()
    
    if nav_about and not is_generating:
        if st.session_state.get("current_page") != "about":
            st.session_state["current_page"] = "about"
            try:
                st.query_params.page = "about"
            except:
                pass
            st.rerun()
    
    # Display content based on selected page - completely separate pages
    if current_page == "assistant":
        research_assistant_page(ingestion, query_pipeline)
    elif current_page == "report":
        research_report_page()
    elif current_page == "about":
        research_about_page()
    
    # Sidebar: Documents list below navigation
    st.sidebar.divider()
    sidebar_documents()


if __name__ == "__main__":
    main()


